import re
import tiktoken
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from docx import Document as DocDocument
import logging
from langchain.prompts import ChatPromptTemplate
from operator import itemgetter
from langchain_openai import AzureChatOpenAI
import os
import sys
from langchain.schema import StrOutputParser
import pickle
import pandas
import openpyxl
import uuid
import json
from dotenv import load_dotenv
import logging
from tenacity import retry, stop_after_attempt, wait_fixed, before_log
from azure.core.credentials import AzureKeyCredential
from langchain.docstore.document import Document
from langchain_openai import AzureOpenAIEmbeddings
from langchain_community.vectorstores import AzureSearch
from azure.search.documents.indexes.models import (
    FreshnessScoringFunction,
    FreshnessScoringParameters,
    ScoringProfile,
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SimpleField,
    TextWeights,
    SemanticConfiguration,
    SemanticPrioritizedFields,
    SemanticField
)
from datetime import datetime
from azure.identity import DefaultAzureCredential
from langchain_core.exceptions import LangChainException # Import this to catch specifically
from collections import defaultdict

# Configure logging if not already done globally
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

#logger = logging.getLogger(__name__)
load_dotenv()

aoai_embeddings = AzureOpenAIEmbeddings(
    azure_deployment=os.getenv('AZURE_DEPLOYMENT_EMBEDDINGS'),
    openai_api_version=os.getenv('OPENAI_API_VERSION'),
    azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT_EMBED'),
    openai_api_key = os.getenv('AZURE_OPENAI_API_KEY_EMBED')

)

llm = AzureChatOpenAI(
    api_version="2024-02-15-preview",
    deployment_name="gpt-4o",
    temperature=0.3
)


vector_store_address: str = os.getenv('AZURE_SEARCH_ENDPOINT')

vector_store_password: str = os.getenv(
    'AZURE_SEARCH_ADMIN_KEY')
credential = DefaultAzureCredential()

def get_vector_store_fields() -> any:
    """Code to define the fields for document chumked in vector store
        id, content,source , content_vector and metadata
    """

    fields = [
        SimpleField(
            name="id",
            type=SearchFieldDataType.String,
            key=True,
            filterable=True,
        ),
        SearchableField(
            name="content",
            type=SearchFieldDataType.String,
            searchable=True,
            filterable=True,
        ),
        SearchableField(
            name="source",
            type=SearchFieldDataType.String,
            searchable=True,
            filterable=True,
        ),
        SearchField(
            name="content_vector",
            type=SearchFieldDataType.Collection(
                SearchFieldDataType.Single),
            searchable=True,
            vector_search_dimensions=len(
                aoai_embeddings.embed_query("Text")),
            vector_search_profile_name="myExhaustiveKnnProfile",
        ),
        SearchableField(
            name="metadata",
            type=SearchFieldDataType.String,
            searchable=True,
            filterable=True,
        ),
    ]
    return fields


def get_vector_store_client(vector_store_index) -> any:
    '''Get AI Search client'''
    index_name: str = vector_store_index
    # Adding a custom scoring profile with a freshness function
    sc_name = "csrd_scoring_profile"
    sc = ScoringProfile(
        name=sc_name,
        text_weights=TextWeights(weights={
            "content":10,
            "content_vector": 8
        }),
        function_aggregation="sum"
    )

    semantic_configuration_name = 'csrd_semantic_configuration'
    semantic_config = SemanticConfiguration(
        name=semantic_configuration_name,
        prioritized_fields=SemanticPrioritizedFields(
            title_field=None,
            content_fields=[SemanticField(field_name='content')],
            keywords_fields=None,
        )
    )

    vector_store: AzureSearch = AzureSearch(
        search_type='semantic_hybrid',
        scoring_profiles=[sc],
        default_scoring_profile=sc_name,
        semantic_configurations=[semantic_config],
        semantic_configuration_name=semantic_configuration_name,
        azure_search_endpoint=vector_store_address,
        azure_search_key=None,
        index_name=index_name,
        embedding_function=aoai_embeddings.embed_query,
        fields=get_vector_store_fields(),
        azure_search_credential=credential
    )
    return vector_store


def store_document_chunks_vector_store(document_name, document_chunks,
                                         vector_store_index):
    '''Generate the vector embeddings for chunk and persist chunks along with
      embeddings in AI search index.
    '''
    try:
        logging.info(
            'Chunk embedding and vector store persistence started at : %s',
            datetime.now())

        vector_store = get_vector_store_client(
            vector_store_index=vector_store_index)
        vector_store.add_documents(documents=document_chunks)
        logging.info(
            'Chunk embedding and vector store persistence finished at : %s',
            datetime.now())

    except LangChainException as lc_error: # Catch LangChainException specifically
        logging.error(f"LangChainException caught during document indexing: {lc_error}")
        # The error object 'lc_error' directly contains the list of IndexingResult objects
        if isinstance(lc_error.args[0], list): # Check if the first arg is a list (of IndexingResult)
            for result in lc_error.args[0]:
                if not result.succeeded:
                    logging.error(f"Document ID: {result.key}, Error Message: {result.error_message}")
        exc_info = sys.exc_info()
        logging.error(
            'Error occurred while persisting document chunks for document: %s to vector store index : %s'
            + ': Error: %s', document_name, vector_store_index, exc_info)
        raise RuntimeError(
             'Error occurred while persisting document chunks for document: %s to vector store index : %s'
             + ': Error: %s', document_name, vector_store_index,
             exc_info) from lc_error # Chain the exception
    except RuntimeError as runtime_error:
        exc_info = sys.exc_info()
        logging.error(
            'Error occurred while persisting document chunks for document: %s to vector store index : %s'
            + ': Error: %s', document_name, vector_store_index, exc_info)
        raise RuntimeError(
             'Error occurred while persisting document chunks for document: %s to vector store index : %s'
             + ': Error: %s', document_name, vector_store_index,
             exc_info) from runtime_error

@retry(stop=stop_after_attempt(3), wait=wait_fixed(2))
def get_relevant_documents(vector_store_index, query, documents_to_consider) -> any:
    '''Custom function to execute AI search based on given query and filter
    based on source documents'''
    try:
        logging.info(
            'get_relevant_document() function execution started at : %s',
            datetime.now())
        # Initialize the vector store client with the specified index
        vector_store = get_vector_store_client(vector_store_index
                                               =vector_store_index)

        def get_filter_params(files):
            '''Inner function for generating search filter params string for
            field "source".'''
            return "|".join(file for file in files)

        '''Fetch the relevant document utulizing the semantic hybrid search
          with semantic reranker score.'''

        '''Convert Doc string to list'''
        docs_list = documents_to_consider.split(",")

        filter_param = get_filter_params(docs_list)
        search_results = vector_store.semantic_hybrid_search_with_score_and_rerank(
            query=query,
            k=3, #Put as global var afterwards
            filters=f"search.in(source,'{filter_param}','|')"
        )

        results = []
        for search_result in search_results:
            #print(search_result)
            document = search_result[0]

            document.metadata['search_score'] = search_result[1]
            document.metadata['reranker_score'] = search_result[2]
            results.append(document)

        logging.info('get_relevant_document() function execution finished at : %s'
                     , datetime.now())
        if len(results) == 0:
            document = Document(page_content='No context avaiable in document')
            document.metadata = {
            'page': 0,
            'source': '',
            'search_score': 0,
            'reranker_score': 0
            }
            results.append(document)

        return results

    except RuntimeError as runtime_error:
        exc_info = sys.exc_info()
        logging.error(
            'Error occurred while executing AI search query : %s to vector'
            + ' store index : %s : Error: %s', query, vector_store_index
            , exc_info)
        raise RuntimeError(
             'Error occurred while executing AI search query : %s to vector store index : %s'
             + ': Error: %s', query, vector_store_index,
             exc_info) from runtime_error


logger = logging.getLogger(__name__)

@dataclass
class ChunkMetadata:
    """Metadata for document chunks"""
    chunk_id: str
    paragraph_indices: List[int]
    word_count: int
    token_count: int
    starts_with_heading: bool
    contains_list: bool
    formatting_info: Dict

@dataclass
class DocumentChunk:
    """Enhanced document chunk with metadata"""
    content: str
    metadata: ChunkMetadata
    # original_paragraphs: List[str]

class EnhancedDocumentChunker:
    """Enhanced chunking system that maintains paragraph integrity"""

    def __init__(self, max_tokens: int = 2000, model_name: str = "o200k_base"):
        self.max_tokens = max_tokens
        self.tokenizer = tiktoken.get_encoding(model_name)
        self.min_chunk_tokens = max_tokens // 4  # Minimum 25% of max tokens
        self.min_paragraph_chars = 50

    def extract_paragraphs_from_docx(self, doc_path: str) -> List[Dict]:
        """Extract paragraphs from Word document with metadata"""
        doc = DocDocument(doc_path)
        paragraphs = []

        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip() or len(para.text.strip()) < self.min_paragraph_chars:
                continue


            # Extract paragraph metadata
            metadata = {
                'index': i,
                'style': para.style.name if para.style else 'Normal',
                'alignment': str(para.alignment) if para.alignment else 'LEFT',
                'is_heading': 'Heading' in (para.style.name if para.style else ''),
                'is_list_item': para.style.name and 'List' in para.style.name,
                'formatting': self._extract_formatting(para),
                'level': self._get_heading_level(para)
            }

            paragraphs.append({
                'text': para.text.strip(),
                'metadata': metadata,
                'word_count': len(para.text.split()),
                'token_count': len(self.tokenizer.encode(para.text))
            })

        return paragraphs

    def _extract_formatting(self, paragraph) -> Dict:
        """Extract formatting information from paragraph"""
        formatting = {
            'bold': False,
            'italic': False,
            'underline': False,
            'font_sizes': set(),
            'font_names': set()
        }

        for run in paragraph.runs:
            if run.bold:
                formatting['bold'] = True
            if run.italic:
                formatting['italic'] = True
            if run.underline:
                formatting['underline'] = True
            if run.font.size:
                formatting['font_sizes'].add(run.font.size.pt)
            if run.font.name:
                formatting['font_names'].add(run.font.name)

        # Convert sets to lists for JSON serialization
        formatting['font_sizes'] = list(formatting['font_sizes'])
        formatting['font_names'] = list(formatting['font_names'])

        return formatting

    def _get_heading_level(self, paragraph) -> int:
        """Extract heading level from paragraph style"""
        if not paragraph.style:
            return 0

        style_name = paragraph.style.name
        if 'Heading' in style_name:
            # Extract number from "Heading 1", "Heading 2", etc.
            match = re.search(r'Heading (\d+)', style_name)
            if match:
                return int(match.group(1))

        return 0

    # def chunk_by_semantic_boundaries(self, paragraphs: List[Dict]) -> List[DocumentChunk]:
    #     """Create chunks respecting semantic boundaries"""
    #     chunks = []
    #     current_chunk_paras = []
    #     current_tokens = 0
    #     chunk_counter = 0

    #     for i, para in enumerate(paragraphs):
    #         para_tokens = para['token_count']

    #         # Check if we should start a new chunk
    #         should_start_new = self._should_start_new_chunk(
    #             current_tokens, para_tokens, para, current_chunk_paras
    #         )

    #         if should_start_new and current_chunk_paras:
    #             # Create chunk from current paragraphs
    #             chunk = self._create_chunk(current_chunk_paras, chunk_counter)
    #             chunks.append(chunk)
    #             chunk_counter += 1

    #             # Start new chunk
    #             current_chunk_paras = [para]
    #             current_tokens = para_tokens
    #         else:
    #             # Add to current chunk
    #             current_chunk_paras.append(para)
    #             current_tokens += para_tokens

    #     # Handle remaining paragraphs
    #     if current_chunk_paras:
    #         chunk = self._create_chunk(current_chunk_paras, chunk_counter)
    #         chunks.append(chunk)

    #     return chunks

    def _should_start_new_chunk(self, current_tokens: int, para_tokens: int,
                               para: Dict, current_chunk_paras: List[Dict]) -> bool:
        """Determine if we should start a new chunk"""
        # Don't start new chunk if current chunk is too small
        if current_tokens < self.min_chunk_tokens:
            return False

        # Always start new chunk if adding paragraph exceeds max tokens
        if current_tokens + para_tokens > self.max_tokens:
            return True

        # Start new chunk at major headings (Heading 1, Heading 2)
        if para['metadata']['is_heading'] and para['metadata']['level'] <= 2:
            return True

        # Start new chunk if there's a significant style change
        if self._is_significant_style_change(para, current_chunk_paras):
            return True

        return False

    def _is_significant_style_change(self, para: Dict, current_chunk_paras: List[Dict]) -> bool:
        """Check for significant style changes that warrant new chunk"""
        if not current_chunk_paras:
            return False

        last_para = current_chunk_paras[-1]

        # Check for transition from/to list items
        if para['metadata']['is_list_item'] != last_para['metadata']['is_list_item']:
            return True

        # Check for heading level changes
        if (para['metadata']['level'] > 0 and
            last_para['metadata']['level'] > 0 and
            abs(para['metadata']['level'] - last_para['metadata']['level']) > 1):
            return True

        return False

    def _create_chunk(self, paragraphs: List[Dict], chunk_id: int) -> DocumentChunk:
        """Create a DocumentChunk from paragraphs"""
        content = '\n\n'.join(para['text'] for para in paragraphs)
        # original_paragraphs = [para['text'] for para in paragraphs]

        # Calculate metadata
        total_tokens = sum(para['token_count'] for para in paragraphs)
        total_words = sum(para['word_count'] for para in paragraphs)
        paragraph_indices = [para['metadata']['index'] for para in paragraphs]

        starts_with_heading = paragraphs[0]['metadata']['is_heading'] if paragraphs else False
        contains_list = any(para['metadata']['is_list_item'] for para in paragraphs)

        # Aggregate formatting info
        formatting_info = {
            'styles': list(set(para['metadata']['style'] for para in paragraphs)),
            'heading_levels': list(set(para['metadata']['level'] for para in paragraphs if para['metadata']['level'] > 0)),
            'has_formatting': any(
                any(para['metadata']['formatting'].values()) for para in paragraphs
            )
        }

        metadata = ChunkMetadata(
            chunk_id=f"chunk_{chunk_id}",
            paragraph_indices=paragraph_indices,
            word_count=total_words,
            token_count=total_tokens,
            starts_with_heading=starts_with_heading,
            contains_list=contains_list,
            formatting_info=formatting_info
        )

        return DocumentChunk(
            content=content,
            metadata=metadata,
            # original_paragraphs=original_paragraphs
        )

    def chunk_with_overlap(self, paragraphs: List[Dict], overlap_sentences: int = 2) -> List[DocumentChunk]:
        """Create chunks with sentence-level overlap for better context"""
        chunks = []
        current_chunk_paras = []
        current_tokens = 0
        chunk_counter = 0

        for i, para in enumerate(paragraphs):
            para_tokens = para['token_count']

            if current_tokens + para_tokens > self.max_tokens and current_chunk_paras:
                # Create chunk with overlap
                chunk = self._create_chunk_with_overlap(
                    current_chunk_paras, chunk_counter, overlap_sentences
                )
                chunks.append(chunk)
                chunk_counter += 1

                # Start new chunk with overlap from previous
                overlap_paras = self._get_overlap_paragraphs(current_chunk_paras, overlap_sentences)
                current_chunk_paras = overlap_paras + [para]
                current_tokens = sum(p['token_count'] for p in current_chunk_paras)
            else:
                current_chunk_paras.append(para)
                current_tokens += para_tokens

        # Handle remaining paragraphs
        if current_chunk_paras:
            chunk = self._create_chunk_with_overlap(
                current_chunk_paras, chunk_counter, overlap_sentences
            )
            chunks.append(chunk)

        return chunks

    def _create_chunk_with_overlap(self, paragraphs: List[Dict], chunk_id: int, overlap_sentences: int) -> DocumentChunk:
        """Create chunk with overlap handling"""
        return self._create_chunk(paragraphs, chunk_id)

    def _get_overlap_paragraphs(self, paragraphs: List[Dict], overlap_sentences: int) -> List[Dict]:
        """Get paragraphs for overlap"""
        if not paragraphs or overlap_sentences <= 0:
            return []

        # Take last few paragraphs for overlap
        overlap_paras = paragraphs[-min(overlap_sentences, len(paragraphs)):]
        return overlap_paras

    def chunk_by_structure(self, paragraphs: List[Dict]) -> List[DocumentChunk]:
        """Create chunks based on document structure (sections, headings)"""
        chunks = []
        current_section = []
        current_tokens = 0
        chunk_counter = 0

        for para in paragraphs:
            # Start new section at major headings
            if para['metadata']['is_heading'] and para['metadata']['level'] <= 2:
                if current_section:
                    # Close current section
                    chunk = self._create_chunk(current_section, chunk_counter)
                    chunks.append(chunk)
                    chunk_counter += 1
                    current_section = []
                    current_tokens = 0

            # Add paragraph to current section
            current_section.append(para)
            current_tokens += para['token_count']

            # Split section if it gets too large
            if current_tokens > self.max_tokens:
                # Find good split point
                split_point = self._find_split_point(current_section)
                if split_point > 0:
                    # Create chunk from first part
                    chunk = self._create_chunk(current_section[:split_point], chunk_counter)
                    chunks.append(chunk)
                    chunk_counter += 1

                    # Continue with remaining paragraphs
                    current_section = current_section[split_point:]
                    current_tokens = sum(p['token_count'] for p in current_section)

        # Handle remaining section
        if current_section:
            chunk = self._create_chunk(current_section, chunk_counter)
            chunks.append(chunk)

        return chunks

    def _find_split_point(self, paragraphs: List[Dict]) -> int:
        """Find optimal split point within a section"""
        if len(paragraphs) <= 1:
            return 0

        # Prefer splitting at sub-headings
        for i in range(len(paragraphs) - 1, 0, -1):
            if paragraphs[i]['metadata']['is_heading'] and paragraphs[i]['metadata']['level'] > 2:
                return i

        # Otherwise split at middle
        return len(paragraphs) // 2

# Usage example and comparison
def demonstrate_chunking():
    """Demonstrate different chunking approaches"""

    # Initialize chunker
    chunker = EnhancedDocumentChunker(max_tokens=2000)

    # Example document processing
    doc_path = "tmp/Reference2.docx"

    try:
        # Extract paragraphs
        paragraphs = chunker.extract_paragraphs_from_docx(doc_path)
        print(f"Extracted {len(paragraphs)} paragraphs")

        # # Method 1: Semantic boundary chunking
        # semantic_chunks = chunker.chunk_by_semantic_boundaries(paragraphs)
        # print(f"Semantic chunking: {len(semantic_chunks)} chunks")

        # # Method 2: Structure-based chunking
        # structure_chunks = chunker.chunk_by_structure(paragraphs)
        # print(f"Structure chunking: {len(structure_chunks)} chunks")

        # Method 3: Overlap chunking
        overlap_chunks = chunker.chunk_with_overlap(paragraphs, overlap_sentences=2)
        print(f"Overlap chunking: {len(overlap_chunks)} chunks")

        langchain_docs = []
        for i, chunk in enumerate(overlap_chunks):
            # For the purpose of demonstration and immediate use,
            # using a simple ID. In a real scenario, you'd want a more robust unique ID.
            doc_id = f"doc_{i}_{chunk.metadata.chunk_id}"
            langchain_docs.append(Document(
                page_content=chunk.content,
                metadata={
                    
                    "source": 'Reference2.docx',
                    "chunk_id": chunk.metadata.chunk_id,
                    "paragraph_indices": chunk.metadata.paragraph_indices,
                    "word_count": chunk.metadata.word_count,
                    "token_count": chunk.metadata.token_count,
                    "starts_with_heading": chunk.metadata.starts_with_heading,
                    "contains_list": chunk.metadata.contains_list,
                    "formatting_info": str(chunk.metadata.formatting_info) # Convert dict to string for metadata
                }
            ))
        
        return langchain_docs

        # Display chunk information
        # for i, chunk in enumerate(semantic_chunks[:3]):  # Show first 3 chunks
        #     print(f"\nChunk {i+1}:")
        #     print(f"Tokens: {chunk.metadata.token_count}")
        #     print(f"Paragraphs: {len(chunk.metadata.paragraph_indices)}")
        #     print(f"Starts with heading: {chunk.metadata.starts_with_heading}")
        #     print(f"Contains list: {chunk.metadata.contains_list}")
        #     print(f"Content preview: {chunk.content[:200]}...")


    except Exception as e:
        print(f"Error processing document: {e}")

# Improved version of your original functions
def enhanced_split_text_into_chunks(text: str, max_tokens: int = 2000) -> List[str]:
    """Enhanced text splitting that respects sentence boundaries"""
    tokenizer = tiktoken.get_encoding("o200k_base")

    # Better sentence splitting with multiple delimiters
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)

    chunks = []
    current_chunk = []
    current_tokens = 0

    for sentence in sentences:
        sentence_tokens = len(tokenizer.encode(sentence))

        # If single sentence exceeds max tokens, split it further
        if sentence_tokens > max_tokens:
            # Split long sentence by clauses
            clauses = re.split(r'[,;:]', sentence)
            for clause in clauses:
                clause_tokens = len(tokenizer.encode(clause))
                if current_tokens + clause_tokens > max_tokens and current_chunk:
                    chunks.append(' '.join(current_chunk))
                    current_chunk = [clause.strip()]
                    current_tokens = clause_tokens
                else:
                    current_chunk.append(clause.strip())
                    current_tokens += clause_tokens
        else:
            if current_tokens + sentence_tokens > max_tokens and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_tokens = sentence_tokens
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

def enhanced_chunk_paragraphs_by_limit(text: str, max_tokens: int = 2000,
                                     paragraph_delimiter: str = "----$----") -> List[str]:
    """Enhanced paragraph chunking with better boundary detection"""
    tokenizer = tiktoken.get_encoding("cl100k_base")

    # Split by paragraph delimiter
    raw_paragraphs = [p.strip() for p in text.split(paragraph_delimiter) if p.strip()]

    chunks = []
    current_chunk = []
    current_tokens = 0

    for para in raw_paragraphs:
        para_tokens = len(tokenizer.encode(para))

        # Handle oversized paragraphs
        if para_tokens > max_tokens:
            # If current chunk has content, close it
            if current_chunk:
                chunks.append(paragraph_delimiter + "\n" + f"\n{paragraph_delimiter}\n".join(current_chunk) + "\n" + paragraph_delimiter)
                current_chunk = []
                current_tokens = 0

            # Split oversized paragraph
            sub_chunks = enhanced_split_text_into_chunks(para, max_tokens)
            for sub_chunk in sub_chunks:
                chunks.append(paragraph_delimiter + "\n" + sub_chunk + "\n" + paragraph_delimiter)
        else:
            # Check if adding paragraph exceeds limit
            if current_tokens + para_tokens > max_tokens and current_chunk:
                chunks.append(paragraph_delimiter + "\n" + f"\n{paragraph_delimiter}\n".join(current_chunk) + "\n" + paragraph_delimiter)
                current_chunk = [para]
                current_tokens = para_tokens
            else:
                current_chunk.append(para)
                current_tokens += para_tokens

    # Add remaining paragraphs
    if current_chunk:
        chunks.append(paragraph_delimiter + "\n" + f"\n{paragraph_delimiter}\n".join(current_chunk) + "\n" + paragraph_delimiter)

    return chunks



def get_combined_query_from_diclosure_config(disclosure_config) -> str:
        """
        Generates a combined query from the disclosure configuration.
        """
        try:
            prompt_template = """
                System: You are an optimized and concise query generator used for keyword, semantic, and vector search for best results. Your task is to generate a concise query from the given question and user instructions which outline the expected details in the answer.
                The generation of the final query should be completely based on provided question and user instructions data only.
                Question: {question}
                User Instructions: {user_instructions}
                Answer:
            """
            prompt = ChatPromptTemplate.from_template(prompt_template)
            
            rag_chain_query = (
                {
                    "question": itemgetter("question"),
                    "user_instructions": itemgetter("user_instructions")
                }
                | prompt
                |llm
                | StrOutputParser()
            )

            result = rag_chain_query.invoke({
                "question": str(disclosure_config['kpi_query']),
                "user_instructions": str(disclosure_config['prompt_instructions'])
            })

            # logger.info("Combined query generated by LLM: %s", result)
            return result
        except Exception as ke:
            logger.error("Exception occured in get_combined_query_from_diclosure_config: %s", str(ke), exc_info=True)
            raise 


# def get_llm_context(disclosure_config, old_context, new_context) -> str:
#         """
#         Generates a combined query from the disclosure configuration.
#         """
#         try:
#             prompt_template = """
#                 **Task:** Document Amendment Assistant
#                 **Objective:** Process multiple paragraphs simultaneously, updating only those that require changes based on new document context.

#                 **Instructions:**
#                 1. Compare each original paragraph with the new document context.
#                 2. For paragraphs requiring updates:
#                     - Seamlessly **blend or rephrase** existing content using the new context.
#                     - If the new context clearly changes the status of a process (e.g., from planning to completion), you may reword or replace earlier planning language accordingly, rather than inserting a separate status sentence.
#                     - Avoid mixing outdated future intent with updated completion facts â€” rewrite the sentence to reflect the new state clearly.
#                     - Do **not** simply append or prepend disjointed sentences.
#                     - Preserve the original **tone**, **style**, and **paragraph structure**. Slight variations in line count are acceptable for clarity.
#                 3. For paragraphs with no relevant updates: do **not** include them in the final output.
#                 4. Maintain overall **meaning, coherence**, and **natural flow** in all modifications.
#                 5. Do **not modify** standalone **headings or subheadings**. If a paragraph appears to be a title (e.g., fewer than 12 words, all-caps, or capitalized like a section heading), do **not** include it in the final output.

#                 ---
#                 **Input Data:**
#                 **Original Paragraphs Array:**
#                 {old_chunks}

#                 **New Document Context (Amendments):**
#                 {new_chunks}

#                 ---
#                 **Processing Rules:**
#                 - Modify only if the new context provides relevant, factual updates.
#                 - Blend new information **within** existing phrasing. You may **reword** or **replace** parts to reflect changes clearly.
#                 - Avoid abrupt insertions or duplications.
#                 - Try to retain paragraph rhythm. Minor structure adjustments are acceptable when needed for smooth integration.
#                 - Do not remove original context unless it's outdated or contradicts the new facts.
#                 - Keep formatting clean and consistent.
#                 - Do **not** modify headings or subheadings â€” skip them.

#                 **Output Format:**
#                 Return a strict JSON array containing **only** the modified paragraphs. Each object in the array should represent a modified paragraph and must include its "paragraph", "content", and a "modification_reason". Do not include paragraphs that did not require modification.

#                 ```json
#                 [
#                   {{
#                     "paragraph": Provided Original Paragraph,
#                     "content": "[Modified content matching original structure]",
#                     "modification_reason": "[Brief explanation of the modification]"
#                   }},
#                   {{
#                     "paragraph": Provided Original Paragraph,
#                     "content": "[Modified content here]",
#                     "modification_reason": "[Reason for modification]"
#                   }}
#                 ]
#                 ```
# """
 
#             prompt = ChatPromptTemplate.from_template(prompt_template)
            
#             rag_chain_query = (
#                 {
#                     # "kpi_query": itemgetter("kpi_query"),
#                     # "prompt_instructions": itemgetter("prompt_instructions"),
#                     "old_chunks": itemgetter("old_context"),
#                     "new_chunks": itemgetter("new_context")
#                 }
#                 | prompt
#                 |llm
#                 | StrOutputParser()
#             )


#             result = rag_chain_query.invoke({
#                 "kpi_query": str(disclosure_config['kpi_query']),
#                 "prompt_instructions": str(disclosure_config['prompt_instructions']),
#                 "old_context": old_context,
#                 "new_context": new_context
#             })

#             return result
#         except Exception as ke:
#             logger.error("Exception occured in get_combined_query_from_diclosure_config: %s", str(ke), exc_info=True)
#             raise 

def get_llm_context(disclosure_config, old_context, new_context) -> str:
    """
    Generates a combined query from the disclosure configuration,
    aligning context generation with the KPI query and preventing hallucination.
    """
    try:
        prompt_template = """
                **Task:** Document Amendment Assistant
                **Objective:** Process multiple paragraphs simultaneously, updating only those that require changes based on new document context and are relevant to the provided KPI query.

                **Instructions:**
                1. Analyze the **KPI Query** to understand the specific performance indicator or area of focus.
                2. Compare each original paragraph with the new document context.
                3. **Crucially, only consider paragraphs for update if they are directly relevant to the KPI Query.**
                4. For paragraphs requiring updates (and relevant to the KPI):
                    - Seamlessly **blend or rephrase** existing content using the new context.
                    - If the new context clearly changes the status of a process (e.g., from planning to completion), you may reword or replace earlier planning language accordingly, rather than inserting a separate status sentence.
                    - Avoid mixing outdated future intent with updated completion facts â€” rewrite the sentence to reflect the new state clearly.
                    - Do **not** simply append or prepend disjointed sentences.
                    - Preserve the original **tone**, **style**, and **paragraph structure**. Slight variations in line count are acceptable for clarity.
                5. For paragraphs with no relevant updates OR paragraphs **not relevant to the KPI Query**: do **not** include them in the final output.
                6. Maintain overall **meaning, coherence**, and **natural flow** in all modifications.
                7. Do **not modify** standalone **headings or subheadings**. If a paragraph appears to be a title (e.g., fewer than 12 words, all-caps, or capitalized like a section heading), do **not** include it in the final output.
                8. **Strictly adhere to the provided contexts. Do not introduce any information not present in either the "Original Paragraphs Array" or the "New Document Context (Amendments)". Avoid hallucination.**

                ---
                **Input Data:**
                **KPI Query:**
                {kpi_query}

                **Original Paragraphs Array:**
                {old_chunks}

                **New Document Context (Amendments):**
                {new_chunks}

                ---
                **Processing Rules:**
                - Modify only if the new context provides relevant, factual updates **AND** the paragraph is relevant to the KPI Query.
                - Blend new information **within** existing phrasing. You may **reword** or **replace** parts to reflect changes clearly.
                - Avoid abrupt insertions or duplications.
                - Try to retain paragraph rhythm. Minor structure adjustments are acceptable when needed for smooth integration.
                - Do not remove original context unless it's outdated, contradicts the new facts, or is not relevant to the KPI.
                - Keep formatting clean and consistent.
                - Do **not** modify headings or subheadings â€” skip them.
                - **Generate content exclusively from the provided `old_chunks` and `new_chunks`. Do not invent or infer information.**

                **Output Format:**
                Return a strict JSON array containing **only** the modified paragraphs that are relevant to the KPI Query. Each object in the array should represent a modified paragraph and must include its "paragraph", "content", and a "modification_reason". Do not include paragraphs that did not require modification or were not relevant to the KPI Query.

                ```json
                [
                  {{
                    "paragraph": "Provided Original Paragraph",
                    "content": "[Modified content matching original structure]",
                    "modification_reason": "[Brief explanation of the modification, including its relevance to the KPI]"
                  }},
                  {{
                    "paragraph": "Provided Original Paragraph",
                    "content": "[Modified content here]",
                    "modification_reason": "[Reason for modification and KPI relevance]"
                  }}
                ]
                ```
        """
        
        # prompt_template = """
        #         **Task:** Document Amendment Assistant
        #         **Objective:** Process multiple paragraphs simultaneously, updating only those that require changes based on new document context and are relevant to the provided KPI query.

        #         **Instructions:**
        #         1. Analyze the **KPI Query** to understand the specific performance indicator or area of focus.
        #         2. Compare each original paragraph with the new document context.
        #         3. **Crucially, only consider paragraphs for update if they are directly relevant to the KPI Query.**
        #         4. For paragraphs requiring updates (and relevant to the KPI):
        #             - Seamlessly **blend or rephrase** existing content using the new context.
        #             - **Prioritize rewriting existing sentences or clauses to seamlessly embed new information, rather than simply appending new sentences. The goal is to make the updated paragraph read as if it was originally written with the new information in mind.**
        #             - If the new context clearly changes the status of a process (e.g., from planning to completion), you may reword or replace earlier planning language accordingly, rather than inserting a separate status sentence.
        #             - Avoid mixing outdated future intent with updated completion facts â€” rewrite the sentence to reflect the new state clearly.
        #             - Do **not** simply append or prepend disjointed sentences.
        #             - Preserve the original **tone**, **style**, and **paragraph structure**. Slight variations in line count are acceptable for clarity.
        #             - **Ensure that the modified content maintains a singular thematic focus within the paragraph. If the new context introduces a significantly different topic, consider if it truly belongs within the existing paragraph's core subject, or if it might be better omitted based on the KPI relevance.**
        #         5. For paragraphs with no relevant updates OR paragraphs **not relevant to the KPI Query**: do **not** include them in the final output.
        #         6. Maintain overall **meaning, coherence**, and **natural flow** in all modifications.
        #         7. Do **not modify** standalone **headings or subheadings**. If a paragraph appears to be a title (e.g., fewer than 12 words, all-caps, or capitalized like a section heading), do **not** include it in the final output.
        #         8. **Strictly adhere to the provided contexts. Do not introduce any information not present in either the "Original Paragraphs Array" or the "New Document Context (Amendments)". Avoid hallucination.**

        #         ---
        #         **Integration Example:**
        #         To illustrate the desired blending, consider this:
        #         * **Original Paragraph:** "The company invested in leadership training to improve employee communication."
        #         * **New Context:** "Training now includes modules on active listening and conflict resolution."
        #         * **Good Integration (Desired Output):** "The company invested in leadership training, now enhanced with modules on active listening and conflict resolution, to improve employee communication."
        #         * **Poor Integration (AVOID this style):** "The company invested in leadership training to improve employee communication. Additionally, training now includes modules on active listening and conflict resolution."
        #         ---

        #         **Input Data:**
        #         **KPI Query:**
        #         {kpi_query}

        #         **Original Paragraphs Array:**
        #         {old_chunks}

        #         **New Document Context (Amendments):**
        #         {new_chunks}

        #         ---
        #         **Processing Rules:**
        #         - Modify only if the new context provides relevant, factual updates **AND** the paragraph is relevant to the KPI Query.
        #         - Blend new information **within** existing phrasing. You may **reword** or **replace** parts to reflect changes clearly.
        #         - **Focus on integrating new facts into the existing narrative flow, avoiding standalone additions.**
        #         - Avoid abrupt insertions or duplications.
        #         - Try to retain paragraph rhythm. Minor structure adjustments are acceptable when needed for smooth integration.
        #         - Do not remove original context unless it's outdated, contradicts the new facts, or is not relevant to the KPI.
        #         - Keep formatting clean and consistent.
        #         - Do **not** modify headings or subheadings â€” skip them.
        #         - **Generate content exclusively from the provided `old_chunks` and `new_chunks`. Do not invent or infer information.**

        #         **Output Format:**
        #         Return a strict JSON array containing **only** the modified paragraphs that are relevant to the KPI Query. Each object in the array should represent a modified paragraph and must include its "paragraph", "content", and a "modification_reason". Do not include paragraphs that did not require modification or were not relevant to the KPI Query.

        #         ```json
        #         [
        #           {{
        #             "paragraph": "Provided Original Paragraph",
        #             "content": "[Modified content matching original structure]",
        #             "modification_reason": "[Brief explanation of the modification, including its relevance to the KPI]"
        #           }},
        #           {{
        #             "paragraph": "Provided Original Paragraph",
        #             "content": "[Modified content here]",
        #             "modification_reason": "[Reason for modification and KPI relevance]"
        #           }}
        #         ]
        #         ```
        # """
        
        prompt = ChatPromptTemplate.from_template(prompt_template)

        rag_chain_query = (
            {
                "kpi_query": itemgetter("kpi_query"),
                "old_chunks": itemgetter("old_context"),
                "new_chunks": itemgetter("new_context")
            }
            | prompt
            | llm  # Uncomment and ensure 'llm' is defined and imported
            | StrOutputParser()
        )

        result = rag_chain_query.invoke({
            "kpi_query": str(disclosure_config.get('kpi_query', '')),
            "prompt_instructions": str(disclosure_config.get('prompt_instructions', '')),
            "old_context": old_context,
            "new_context": new_context
        })

        return result
    except Exception as ke:
        # Assuming 'logger' is defined and imported
        # logger.error("Exception occurred in get_combined_query_from_diclosure_config: %s", str(ke), exc_info=True)
        raise

def clean_and_parse_json(json_string):
    """
    Cleans a JSON string by removing any leading/trailing backticks or 'json'
    markers and then parses it into a Python object.
    """
    # Remove leading/trailing ```json and ```
    cleaned_string = re.sub(r'```json\s*', '', json_string, flags=re.DOTALL)
    cleaned_string = re.sub(r'\s*```', '', cleaned_string, flags=re.DOTALL)

    try:
        return json.loads(cleaned_string)
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {e}")
        print(f"Problematic JSON string: {cleaned_string}")
        return None # Return None or raise an error as appropriate

# extracted_results = []
# def refresh(disc_code):
#     # Fetch Relevant data of old document for KPI
#     query = get_combined_query_from_diclosure_config(disc_code)

#     chunk_old = get_relevant_documents('og_accenture' ,query,disc_code['docs_to_consider_old'] )
#     chunk_new = get_relevant_documents('og_accenture' ,query,disc_code['docs_to_consider_new'] )
   
#     #Extract paragraphs based on paragraph incides of old in metadata
#     # load old document
#     doc = DocDocument('tmp/Base2.docx')
#     paragraphs_old_doc = doc.paragraphs
#     def chunk_list(data, size):
#         for i in range(0, len(data), size):
#             yield data[i:i + size]
#     for chunk in chunk_old[:1]:  # only one chunk for demonstration (O(n3) optimization is a separate concern)
#         all_paras = []
#         indices = chunk.metadata.get('paragraph_indices')
#         for index in indices:
#             all_paras.append(paragraphs_old_doc[index].text.strip())

#         # Generate the new_chunks_context outside the if/else to avoid duplication
#         new_chunks_context = "\n\n".join(
#             f"{doc.page_content}\n\n)" # Note: The extra ')' might be a typo, check your original prompt format
#             for doc in chunk_new
#             if doc.page_content.strip().lower() != "no context available in document"
#         )

#         # Determine if batching is needed
#         if len(all_paras) > 10:
#             # If paras list is larger than 10, process in batches
#             batch_size = 10
#             for batch_paras in chunk_list(all_paras, batch_size):
#                 llm_raw_output = get_llm_context(
#                     disc_code,
#                     json.dumps(batch_paras),
#                     new_chunks_context
#                 )

#                 print(f"Raw LLM output (batch): {llm_raw_output}")
                
#                 # Clean and parse the JSON string
#                 parsed_response = clean_and_parse_json(llm_raw_output)
#                 if parsed_response:
#                     # Add 'kpi' to each object in parsed_response
#                     for obj in parsed_response:
#                         obj['kpi'] = disc_code['kpi_query']
                    
#                     # Extend the global list with updated objects
#                     extracted_results.extend(parsed_response)

#                     print(f"Added {len(parsed_response)} objects to extracted_results. Total: {len(extracted_results)}")

#                 # if parsed_response:
#                 #     # Extend the global extracted_results array with objects from the current batch
#                 #     extracted_results.extend({'parsed_response':parsed_response, 'kpi':disc_code['kpi_query']})
#                 #     extracted_results.extend(parsed_response)

#                 #     print(f"Added {len(parsed_response)} objects to extracted_results. Total: {len(extracted_results)}")
#                 else:
#                     print("Skipping addition to extracted_results due to JSON parsing error.")

#         else:
#             # If paras list is 10 or less, process as a single call
#             llm_raw_output = get_llm_context(
#                 disc_code,
#                 json.dumps(all_paras),
#                 new_chunks_context
#             )

#             print(f"Raw LLM output (single call): {llm_raw_output}")

#             # Clean and parse the JSON string
#             parsed_response = clean_and_parse_json(llm_raw_output)
#             if parsed_response:
#                 # Add 'kpi' to each object in parsed_response
#                 for obj in parsed_response:
#                     obj['kpi'] = disc_code['kpi_query']
                
#                 # Extend the global list with updated objects
#                 extracted_results.extend(parsed_response)

#                 print(f"Added {len(parsed_response)} objects to extracted_results. Total: {len(extracted_results)}")

#             # if parsed_response:
#             #     # Extend the global extracted_results array with objects from the single call

#             #     extracted_results.extend({'parsed_response':parsed_response, 'kpi':disc_code['kpi_query']})
#             #     print(f"Added {len(parsed_response)} objects to extracted_results. Total: {len(extracted_results)}")
#             else:
#                 print("Skipping addition to extracted_results due to JSON parsing error.")

# # After the loop, extracted_results will contain all the parsed objects
#     print("\nFinal extracted_results:")
#     with open('output_frank.pkl' , 'wb') as fb:
#         pickle.dump(extracted_results , fb)
def call_llm_to_summarize(content: str) -> str:
    """
    Sends repeated paragraph content to the LLM and returns a concise summary.
    """
    try:
        prompt_template = """
        System: You are a professional document summarizer. Your job is to concisely summarize repetitive content extracted from similar paragraphs in an ESG report. 
        Provide a meaningful and clear summary capturing the main idea, ensuring no important point is lost.
        Do not repeat any sentence or information. Keep the summary crisp and non-redundant.

        Content to summarize:
        {content}

        Summary:
        """

        prompt = ChatPromptTemplate.from_template(prompt_template)

        summary_chain = (
            {
                "content": itemgetter("content")
            }
            | prompt
            | llm
            | StrOutputParser()
        )

        result = summary_chain.invoke({
            "content": content
        })

        return result

    except Exception as e:
        logger.error("Exception occurred in call_llm_to_summarize: %s", str(e), exc_info=True)
        raise

extracted_results = []
paragraph_to_objects = defaultdict(list)
paragraph_to_key = defaultdict(list)
summarized_paragraphs = {}           # paragraph â†’ summary
summarized_paragraphs_set = set()    # set of paragraph texts

serials = defaultdict(int)
# To track (kpi_code, paragraph) pairs already added to extracted_results
seen_kpi_paragraph = set()

def refresh(disc_code):
    query = get_combined_query_from_diclosure_config(disc_code)

    chunk_old = get_relevant_documents('og_accenture', query, disc_code['docs_to_consider_old'])
    chunk_new = get_relevant_documents('og_accenture', query, disc_code['docs_to_consider_new'])

    doc = DocDocument('tmp/Base2.docx')
    paragraphs_old_doc = doc.paragraphs

    def chunk_list(data, size):
        for i in range(0, len(data), size):
            yield data[i:i + size]

    for chunk in chunk_old[:1]:  # process first chunk only for demo
        all_paras = []
        indices = chunk.metadata.get('paragraph_indices')
        for index in indices:
            all_paras.append(paragraphs_old_doc[index].text.strip())
        
        new_chunks_context = "\n\n".join(
            f"{doc.page_content}\n\n"
            for doc in chunk_new
            if doc.page_content.strip().lower() != "no context available in document"
        )

        def process_batch(para):

            llm_raw_output = get_llm_context(
                disc_code,
                json.dumps(para),
                new_chunks_context
            )
            print(f"Raw LLM output: {llm_raw_output}")

            parsed_response = clean_and_parse_json(llm_raw_output)
            if parsed_response:
                current_kpi = disc_code['kpi_query']
                current_kpi_code = disc_code['kpi_code']

                for obj in parsed_response:
                    obj['SERIAL'] = current_kpi_code
                    obj['kpi'] = current_kpi
                    paragraph = obj.get("paragraph")
                    if not paragraph:
                        continue

                    # Add object to global tracking
                    paragraph_to_objects[paragraph].append(obj)

                    # Track which KPIs this paragraph appears in
                    if current_kpi_code not in paragraph_to_key[paragraph]:
                        paragraph_to_key[paragraph].append(current_kpi_code)

                # Process all paragraphs
                for paragraph, objs in paragraph_to_objects.items():
                    if len(paragraph_to_key[paragraph]) > 1:
                        # Paragraph is repeated across KPIs
                        if paragraph not in summarized_paragraphs_set:
                            # Summarize only once
                            combined_content = "\n".join(o['content'] for o in objs if 'content' in o)
                            summary = call_llm_to_summarize(combined_content)
                            summarized_paragraphs[paragraph] = summary
                            summarized_paragraphs_set.add(paragraph)
                        else:
                            summary = summarized_paragraphs[paragraph]

                        for obj in objs:
                            obj['content'] = summary
                            obj['also_modified_in'] = paragraph_to_key[paragraph]
                    else:
                        for obj in objs:
                            obj['also_modified_in'] = None

                for paragraph, kpi_codes in paragraph_to_key.items():
                    for kpi_code in kpi_codes:
                        if (kpi_code, paragraph) in seen_kpi_paragraph:
                            continue  # skip duplicates
                        serials[kpi_code] += 1

                        # Find all objects for this paragraph and KPI
                        # They all share the same paragraph and KPI, so we just pick one to modify and add
                        objs_for_kpi = [obj for obj in paragraph_to_objects[paragraph] if obj.get('SERIAL') == kpi_code]

                        if not objs_for_kpi:
                            # Defensive fallback: create a minimal object if none found
                            result_obj = {
                                "paragraph": paragraph,
                                "kpi_code": kpi_code,
                                "serial": serials[kpi_code],
                                "also_modified_in": paragraph_to_key[paragraph] if len(paragraph_to_key[paragraph]) > 1 else None,
                            }
                            extracted_results.append(result_obj)
                        else:
                            # Update and add each object only once
                            # But usually we only want to add one representative object per paragraph-KPI
                            obj = objs_for_kpi[0]
                            obj['serial'] = serials[kpi_code]  # assign unique serial
                            obj['kpi_code'] = kpi_code
                            extracted_results.append(obj)

                        seen_kpi_paragraph.add((kpi_code, paragraph))

                # new_paragraphs_in_batch = set(obj.get("paragraph") for obj in parsed_response if obj.get("paragraph"))

                # for para in new_paragraphs_in_batch:
                #     extracted_results.extend(paragraph_to_objects[para])

                # for objs in paragraph_to_objects.values():
                #     extracted_results.extend(objs)  

                # for i, obj in enumerate(parsed_response):
                #     obj['SERIAL'] = current_kpi_code
                #     obj['kpi']=current_kpi
                #     if obj.get("paragraph") not in paragraph_to_key:
                        
                #         paragraph_to_key[obj.get("paragraph")] = current_kpi_code
                #     else:
                        
                #         obj['also_modified_in'] = paragraph_to_key[obj.get("paragraph")]

                # extracted_results.extend(parsed_response)
                print(f"Added {len(paragraph_to_objects)} objects to extracted_results. Total: {len(extracted_results)}")
            else:
                print("Skipping addition to extracted_results due to JSON parsing error.")

        # Batch processing to handle large paragraphs
        batch_size = 10
        for batch_pairs in chunk_list(all_paras, batch_size):
            process_batch(batch_pairs)

    # Save final output
    print("\nFinal extracted_results:")
    # with open('output_summary_4.pkl', 'wb') as fb:
    #     pickle.dump(extracted_results, fb)

if __name__ == "__main__":
    # Step 1 Create Document Chunks 9Filename and source name p;lelase provide)
    # oc = demonstrate_chunking()
    # Step 2 Store Document Chunks 

    # store_document_chunks_vector_store('testtt.pdf', oc , 'og_accenture')

    # Step 3 KPI  Processing and Retrival

    df = pandas.read_excel('tmp/Refresh_kpis.xlsx' , engine = 'openpyxl')
    row = df.iloc[4]
    for idx , row in df.head(21).iterrows():
        disc_code = {
            'kpi_query': row[11],
            'kpi_code': row[0],
            'prompt_instructions': row[8],
            'docs_to_consider_old': 'Base2.docx',
            'docs_to_consider_new': 'Reference2.docx',
        }
        refresh(disc_code)
